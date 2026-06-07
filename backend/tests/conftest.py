"""tests/conftest.py — Shared pytest fixtures."""
from __future__ import annotations

import io
import pytest
from httpx import ASGITransport, AsyncClient

from app.main import app


@pytest.fixture
def anyio_backend():
    return "asyncio"


@pytest.fixture
async def client():
    async with AsyncClient(
        transport=ASGITransport(app=app),
        base_url="http://test",
    ) as ac:
        yield ac


# ── Minimal fixture files ──────────────────────────────────────────────────────

@pytest.fixture
def sample_pdf_bytes() -> bytes:
    """A tiny single-page PDF with selectable text."""
    # Minimal valid PDF — created inline to avoid binary fixtures
    return (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R"
        b"/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
        b"4 0 obj<</Length 44>>\nstream\n"
        b"BT /F1 12 Tf 100 700 Td (Jane Smith Resume) Tj ET\n"
        b"endstream\nendobj\n"
        b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
        b"xref\n0 6\n0000000000 65535 f\n"
        b"trailer<</Size 6/Root 1 0 R>>\n"
        b"startxref\n0\n%%EOF"
    )


@pytest.fixture
def sample_docx_bytes() -> bytes:
    """A minimal DOCX file created with python-docx in memory."""
    from docx import Document
    doc = Document()
    doc.add_heading("Jane Smith", level=1)
    doc.add_paragraph("Senior Software Engineer")
    doc.add_paragraph("jane@example.com | +1 555 000 1234 | San Francisco, CA")
    doc.add_heading("Experience", level=2)
    doc.add_paragraph("Acme Corp — Lead Engineer (2021-Present)")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
